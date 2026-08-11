from __future__ import annotations

import base64
from pathlib import Path
from typing import Any, cast

import pytest
from docx import Document
from docx.oxml.ns import qn
from typer.testing import CliRunner

from local_document_converter.cli import app
from local_document_converter.domain import (
    DocumentIR,
    DocumentMetadata,
    HeadingBlock,
    ImageBlock,
    ListBlock,
    PageBreakBlock,
    ParagraphBlock,
    SourceInfo,
    TableBlock,
)
from local_document_converter.exceptions import ExportError, OutputExistsError
from local_document_converter.exporters.base import ExportContext
from local_document_converter.exporters.docx_exporter import DocxExporter
from local_document_converter.exporters.registry import ExporterRegistry
from local_document_converter.parsers.markdown import MarkdownParser
from local_document_converter.parsers.registry import ParserRegistry
from local_document_converter.services.conversion_service import (
    ConversionRequest,
    ConversionService,
)

FIXTURES = Path(__file__).parent / "fixtures"
_ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9Zl8sAAAAASUVORK5CYII="
)


def test_docx_export_preserves_semantic_blocks_images_and_warnings(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.md"
    source.write_text("fixture source", encoding="utf-8")
    image_directory = tmp_path / "images"
    image_directory.mkdir()
    (image_directory / "pixel.png").write_bytes(_ONE_PIXEL_PNG)
    destination = tmp_path / "all-blocks.docx"
    context = ExportContext()

    DocxExporter().export(_all_blocks_document(source), destination, context)
    restored = Document(str(destination))

    assert destination.is_file()
    assert [warning.code for warning in context.warnings] == [
        "docx.image_unavailable"
    ]
    assert context.warnings[0].details["uri"] == "images/missing.png"
    assert restored.core_properties.title == "Stage 7 DOCX Sample"
    assert restored.sections[0].page_width == 7772400
    assert restored.sections[0].left_margin == 914400

    paragraphs = {paragraph.text: paragraph for paragraph in restored.paragraphs}
    assert _style_name(paragraphs["Stage 7 DOCX Sample"]) == "Heading 1"
    assert _style_name(paragraphs["Semantic output"]) == "Heading 2"
    assert _style_name(paragraphs["This paragraph verifies Word structure."]) == "Normal"
    for item in ("Bullet one", "Bullet two", "Step one", "Step two"):
        assert cast(Any, paragraphs[item]._p.pPr).numPr is not None

    assert len(restored.tables) == 1
    table = restored.tables[0]
    assert [[cell.text for cell in row.cells] for row in table.rows] == [
        ["Component", "Status"],
        ["Heading", "Ready"],
        ["Table", "Ready"],
    ]
    _assert_table_geometry(table)
    assert len(restored.inline_shapes) == 1
    properties = cast(Any, restored.inline_shapes[0]._inline).docPr
    assert properties.get("descr") == "Embedded sample image"
    assert "[Image unavailable: Missing sample image]" in paragraphs
    assert "Missing image caption" in paragraphs
    assert len(restored.element.xpath(".//w:br[@w:type='page']")) == 1
    assert _style_name(paragraphs["Second page content"]) == "Normal"


def test_markdown_to_docx_uses_atomic_service_and_reports_export_warning(
    tmp_path: Path,
) -> None:
    parsers = ParserRegistry()
    parsers.register(MarkdownParser())
    exporters = ExporterRegistry()
    exporters.register(DocxExporter())
    destination = tmp_path / "output" / "sample.docx"
    service = ConversionService(parsers, exporters)

    result = service.convert(
        ConversionRequest(
            source=FIXTURES / "sample.md",
            output_format="docx",
            destination=destination,
        )
    )

    assert result.destination == destination.resolve()
    assert result.exporter_name == "docx"
    assert result.warning_count == 1
    assert [warning.code for warning in result.warnings] == [
        "docx.image_unavailable"
    ]
    assert Document(str(destination)).tables[0].cell(1, 0).text == "A|B"
    assert not list(destination.parent.glob(".sample.docx.*.tmp"))

    with pytest.raises(OutputExistsError):
        service.convert(
            ConversionRequest(
                source=FIXTURES / "sample.md",
                output_format="docx",
                destination=destination,
            )
        )
    assert Document(str(destination)).tables[0].cell(1, 0).text == "A|B"


def test_cli_converts_markdown_to_reopenable_docx(tmp_path: Path) -> None:
    destination = tmp_path / "含空白 路徑" / "輸出.docx"

    result = CliRunner().invoke(
        app,
        [
            "convert",
            str(FIXTURES / "sample.md"),
            "--to",
            "docx",
            "--output",
            str(destination),
        ],
    )

    assert result.exit_code == 0, result.stderr
    restored = Document(str(destination))
    assert restored.paragraphs[0].text == "範例文件"
    assert _style_name(restored.paragraphs[0]) == "Heading 1"


def test_docx_export_wraps_invalid_destination_and_rejects_options(
    tmp_path: Path,
) -> None:
    document = DocumentIR(source=SourceInfo(path="source.md"))

    with pytest.raises(ExportError, match="could not write DOCX"):
        DocxExporter().export(
            document,
            tmp_path / "missing" / "output.docx",
            ExportContext(),
        )
    with pytest.raises(ExportError, match="unsupported DOCX exporter options"):
        DocxExporter().export(
            document,
            tmp_path / "output.docx",
            ExportContext(options={"unknown": True}),
        )


def test_docx_export_embeds_absolute_windows_image_path(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text("source", encoding="utf-8")
    image_path = tmp_path / "absolute.png"
    image_path.write_bytes(_ONE_PIXEL_PNG)
    destination = tmp_path / "absolute-image.docx"
    context = ExportContext()
    document = DocumentIR(
        source=SourceInfo(path=str(source), media_type="text/markdown"),
        blocks=[
            ImageBlock(
                id="absolute-image",
                order=0,
                uri=str(image_path),
                alt_text="Absolute image",
            )
        ],
    )

    DocxExporter().export(document, destination, context)

    assert len(Document(str(destination)).inline_shapes) == 1
    assert context.warnings == []


def _all_blocks_document(source: Path) -> DocumentIR:
    return DocumentIR(
        source=SourceInfo(path=str(source), media_type="text/markdown"),
        metadata=DocumentMetadata(
            title="Stage 7 DOCX Sample",
            author="Local Document Converter tests",
            page_count=2,
        ),
        blocks=[
            HeadingBlock(
                id="heading-1",
                order=0,
                page_number=1,
                level=1,
                text="Stage 7 DOCX Sample",
            ),
            ParagraphBlock(
                id="paragraph-1",
                order=1,
                page_number=1,
                text="This paragraph verifies Word structure.",
            ),
            HeadingBlock(
                id="heading-2",
                order=2,
                page_number=1,
                level=2,
                text="Semantic output",
            ),
            ListBlock(
                id="bullets",
                order=3,
                page_number=1,
                ordered=False,
                items=["Bullet one", "Bullet two"],
            ),
            ListBlock(
                id="steps",
                order=4,
                page_number=1,
                ordered=True,
                items=["Step one", "Step two"],
            ),
            TableBlock(
                id="table",
                order=5,
                page_number=1,
                caption="Component status",
                column_names=["Component", "Status"],
                rows=[["Heading", "Ready"], ["Table", "Ready"]],
            ),
            ImageBlock(
                id="embedded-image",
                order=6,
                page_number=1,
                uri="images/pixel.png",
                alt_text="Embedded sample image",
                caption="Embedded image caption",
            ),
            ImageBlock(
                id="missing-image",
                order=7,
                page_number=1,
                uri="images/missing.png",
                alt_text="Missing sample image",
                caption="Missing image caption",
            ),
            PageBreakBlock(id="page-break", order=8, page_number=1),
            ParagraphBlock(
                id="paragraph-2",
                order=9,
                page_number=2,
                text="Second page content",
            ),
        ],
    )


def _assert_table_geometry(table: Any) -> None:
    table_xml = cast(Any, table._tbl)
    properties = table_xml.tblPr
    assert properties.first_child_found_in("w:tblW").get(qn("w:w")) == "9360"
    assert properties.first_child_found_in("w:tblInd").get(qn("w:w")) == "120"
    grid_widths = [int(column.get(qn("w:w"))) for column in table_xml.tblGrid]
    assert sum(grid_widths) == 9360
    for row in table.rows:
        cell_widths = [
            int(cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW").get(qn("w:w")))
            for cell in row.cells
        ]
        assert cell_widths == grid_widths


def _style_name(paragraph: Any) -> str:
    assert paragraph.style is not None
    return str(paragraph.style.name)
