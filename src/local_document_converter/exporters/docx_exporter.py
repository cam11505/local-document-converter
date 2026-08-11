"""Semantic, deterministic-structure DOCX export implemented with python-docx."""

from __future__ import annotations

from collections.abc import Sequence
from pathlib import Path
from typing import Any, cast
from urllib.parse import unquote, urlparse

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.image.exceptions import UnrecognizedImageError
from docx.image.image import Image as DocxImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph

from local_document_converter.domain.models import (
    DocumentIR,
    DocumentWarning,
    HeadingBlock,
    ImageBlock,
    ListBlock,
    PageBreakBlock,
    ParagraphBlock,
    TableBlock,
)
from local_document_converter.exceptions import ExportError
from local_document_converter.exporters.base import ExportContext, ExporterCapability

_CONTENT_WIDTH_DXA = 9360
_TABLE_INDENT_DXA = 120
_CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
_HEADER_FILL = "F2F4F7"
_HEADING_BLUE = "2E74B5"
_HEADING_DARK_BLUE = "1F4D78"
_MUTED_TEXT = "666666"


class DocxExporter:
    """Export blocks using the standard_business_brief design preset."""

    capability = ExporterCapability(format_name="docx", output_extension=".docx")

    def export(self, document: DocumentIR, destination: Path, context: ExportContext) -> None:
        if context.options:
            unsupported = ", ".join(sorted(context.options))
            raise ExportError(f"unsupported DOCX exporter options: {unsupported}")

        word_document = Document()
        self._configure_document(word_document, document)
        for block in document.blocks:
            if isinstance(block, HeadingBlock):
                word_document.add_heading(block.text, level=block.level)
            elif isinstance(block, ParagraphBlock):
                word_document.add_paragraph(block.text, style="Normal")
            elif isinstance(block, ListBlock):
                self._add_list(word_document, block)
            elif isinstance(block, TableBlock):
                self._add_table(word_document, block)
            elif isinstance(block, ImageBlock):
                self._add_image(word_document, document, block, context)
            elif isinstance(block, PageBreakBlock):
                word_document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        try:
            word_document.save(str(destination))
        except (OSError, TypeError, ValueError) as exc:
            raise ExportError(f"could not write DOCX output: {destination}") from exc

    @classmethod
    def _configure_document(
        cls, word_document: DocumentObject, document: DocumentIR
    ) -> None:
        section = word_document.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        cls._configure_styles(word_document)
        properties = word_document.core_properties
        properties.title = document.metadata.title or Path(document.source.path).stem
        properties.author = document.metadata.author or "Local Document Converter"
        properties.subject = "Exported from DocumentIR"
        properties.keywords = "DocumentIR, local document conversion"

    @classmethod
    def _configure_styles(cls, word_document: DocumentObject) -> None:
        styles = word_document.styles
        normal = styles["Normal"]
        cls._set_style_font(normal, "Calibri", 11, color="000000")
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.1

        heading_tokens = {
            1: (16, _HEADING_BLUE, 16, 8),
            2: (13, _HEADING_BLUE, 12, 6),
            3: (12, _HEADING_DARK_BLUE, 8, 4),
            4: (11, _HEADING_DARK_BLUE, 6, 3),
            5: (11, "000000", 4, 2),
            6: (10, _MUTED_TEXT, 4, 2),
        }
        for level, (size, color, before, after) in heading_tokens.items():
            style = styles[f"Heading {level}"]
            cls._set_style_font(style, "Calibri", size, bold=True, color=color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

        caption = styles["Caption"]
        cls._set_style_font(caption, "Calibri", 9, italic=True, color=_MUTED_TEXT)
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.space_after = Pt(6)
        caption.paragraph_format.keep_with_next = True

        if "LDC Image Placeholder" not in styles:
            placeholder = styles.add_style("LDC Image Placeholder", WD_STYLE_TYPE.PARAGRAPH)
            placeholder.base_style = caption
        placeholder = styles["LDC Image Placeholder"]
        cls._set_style_font(placeholder, "Calibri", 9, italic=True, color=_MUTED_TEXT)
        placeholder.paragraph_format.space_before = Pt(4)
        placeholder.paragraph_format.space_after = Pt(4)

    @staticmethod
    def _set_style_font(
        style: Any,
        name: str,
        size: int,
        *,
        bold: bool | None = None,
        italic: bool | None = None,
        color: str,
    ) -> None:
        font = style.font
        font.name = name
        font.size = Pt(size)
        font.bold = bold
        font.italic = italic
        font.color.rgb = RGBColor.from_string(color)
        r_pr = style.element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(qn(f"w:{attribute}"), name)

    @classmethod
    def _add_list(cls, word_document: DocumentObject, block: ListBlock) -> None:
        num_id = cls._create_numbering(word_document, ordered=block.ordered)
        for item in block.items:
            paragraph = word_document.add_paragraph(style="Normal")
            paragraph.add_run(item)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.167
            cls._set_paragraph_numbering(paragraph, num_id)

    @classmethod
    def _create_numbering(cls, word_document: DocumentObject, *, ordered: bool) -> int:
        numbering = cast(Any, word_document.part.numbering_part.element)
        abstract_id = cls._next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
        num_id = cls._next_numbering_id(numbering, "w:num", "w:numId")

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi_level = OxmlElement("w:multiLevelType")
        multi_level.set(qn("w:val"), "singleLevel")
        abstract.append(multi_level)

        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_format = OxmlElement("w:numFmt")
        num_format.set(qn("w:val"), "decimal" if ordered else "bullet")
        level.append(num_format)
        level_text = OxmlElement("w:lvlText")
        level_text.set(qn("w:val"), "%1." if ordered else "•")
        level.append(level_text)
        justification = OxmlElement("w:lvlJc")
        justification.set(qn("w:val"), "left")
        level.append(justification)

        paragraph_properties = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        paragraph_properties.append(tabs)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), "720")
        indent.set(qn("w:hanging"), "360")
        paragraph_properties.append(indent)
        level.append(paragraph_properties)
        abstract.append(level)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_reference = OxmlElement("w:abstractNumId")
        abstract_reference.set(qn("w:val"), str(abstract_id))
        num.append(abstract_reference)
        numbering.append(num)
        return num_id

    @staticmethod
    def _next_numbering_id(numbering: Any, tag: str, attribute: str) -> int:
        values = [
            int(value)
            for node in numbering.findall(qn(tag))
            if (value := node.get(qn(attribute))) is not None
        ]
        return max(values, default=0) + 1

    @staticmethod
    def _set_paragraph_numbering(paragraph: Paragraph, num_id: int) -> None:
        properties = paragraph._p.get_or_add_pPr()
        num_properties = OxmlElement("w:numPr")
        level = OxmlElement("w:ilvl")
        level.set(qn("w:val"), "0")
        number = OxmlElement("w:numId")
        number.set(qn("w:val"), str(num_id))
        num_properties.append(level)
        num_properties.append(number)
        properties.append(num_properties)

    @classmethod
    def _add_table(cls, word_document: DocumentObject, block: TableBlock) -> None:
        if block.caption:
            caption = word_document.add_paragraph(block.caption, style="Caption")
            caption.paragraph_format.keep_with_next = True

        width = len(block.column_names or [])
        if width == 0 and block.rows:
            width = len(block.rows[0])
        width = max(width, 1)
        row_count = len(block.rows) + (1 if block.column_names is not None else 0)
        table = word_document.add_table(rows=max(row_count, 1), cols=width)
        table.style = "Table Grid"
        table.autofit = False

        output_rows: list[list[str | None]] = []
        if block.column_names is not None:
            output_rows.append(list(block.column_names))
        output_rows.extend(block.rows)
        if not output_rows:
            output_rows = [[None] * width]

        for row_index, values in enumerate(output_rows):
            row = table.rows[row_index]
            for column_index, value in enumerate(values):
                cell = row.cells[column_index]
                cell.text = value or ""
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cls._format_cell(cell, header=row_index == 0 and block.column_names is not None)
            if row_index == 0 and block.column_names is not None:
                cls._mark_header_row(row)

        column_widths = cls._column_widths(output_rows, width)
        cls._set_table_geometry(table, column_widths)

    @staticmethod
    def _format_cell(cell: _Cell, *, header: bool) -> None:
        if header:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), _HEADER_FILL)
            cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.bold = header
                run_properties = run._element.get_or_add_rPr()
                run_fonts = run_properties.get_or_add_rFonts()
                for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
                    run_fonts.set(qn(f"w:{attribute}"), "Calibri")

    @staticmethod
    def _mark_header_row(row: _Row) -> None:
        properties = row._tr.get_or_add_trPr()
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        properties.append(header)

    @classmethod
    def _column_widths(
        cls, rows: Sequence[Sequence[str | None]], width: int
    ) -> list[int]:
        if width == 1:
            return [_CONTENT_WIDTH_DXA]
        weights = [
            max(4, min(40, max((len(row[index] or "") for row in rows), default=0)))
            for index in range(width)
        ]
        minimum = min(1080, _CONTENT_WIDTH_DXA // width)
        remaining = _CONTENT_WIDTH_DXA - minimum * width
        total_weight = sum(weights)
        result = [minimum + remaining * weight // total_weight for weight in weights]
        result[-1] += _CONTENT_WIDTH_DXA - sum(result)
        return result

    @classmethod
    def _set_table_geometry(cls, table: Table, widths: Sequence[int]) -> None:
        table_xml = cast(Any, table._tbl)
        properties = table_xml.tblPr
        cls._set_or_add_value(properties, "w:tblW", "w:w", _CONTENT_WIDTH_DXA)
        table_width = properties.first_child_found_in("w:tblW")
        table_width.set(qn("w:type"), "dxa")
        cls._set_or_add_value(properties, "w:tblInd", "w:w", _TABLE_INDENT_DXA)
        indent = properties.first_child_found_in("w:tblInd")
        indent.set(qn("w:type"), "dxa")
        layout = properties.first_child_found_in("w:tblLayout")
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            properties.append(layout)
        layout.set(qn("w:type"), "fixed")

        margins = properties.first_child_found_in("w:tblCellMar")
        if margins is None:
            margins = OxmlElement("w:tblCellMar")
            properties.append(margins)
        for side, value in _CELL_MARGINS_DXA.items():
            margin = margins.find(qn(f"w:{side}"))
            if margin is None:
                margin = OxmlElement(f"w:{side}")
                margins.append(margin)
            margin.set(qn("w:w"), str(value))
            margin.set(qn("w:type"), "dxa")

        grid = table_xml.tblGrid
        for column in list(grid):
            grid.remove(column)
        for width in widths:
            column = OxmlElement("w:gridCol")
            column.set(qn("w:w"), str(width))
            grid.append(column)

        for row in table.rows:
            for cell, width in zip(row.cells, widths, strict=True):
                cell_properties = cell._tc.get_or_add_tcPr()
                cls._set_or_add_value(cell_properties, "w:tcW", "w:w", width)
                cell_width = cell_properties.first_child_found_in("w:tcW")
                cell_width.set(qn("w:type"), "dxa")

    @staticmethod
    def _set_or_add_value(parent: Any, tag: str, attribute: str, value: int) -> None:
        element = parent.first_child_found_in(tag)
        if element is None:
            element = OxmlElement(tag)
            parent.append(element)
        element.set(qn(attribute), str(value))

    @classmethod
    def _add_image(
        cls,
        word_document: DocumentObject,
        document: DocumentIR,
        block: ImageBlock,
        context: ExportContext,
    ) -> None:
        image_path, reason = cls._resolve_image(document, block.uri)
        if image_path is None:
            cls._add_image_placeholder(word_document, block)
            context.warnings.append(cls._image_warning(block, reason))
            return

        try:
            image = DocxImage.from_file(str(image_path))
            target_width = min(image.width, Inches(6.5))
            paragraph = word_document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(3)
            shape = paragraph.add_run().add_picture(str(image_path), width=target_width)
            doc_properties = cast(Any, shape._inline).docPr
            doc_properties.set("descr", block.alt_text or image_path.name)
            doc_properties.set("title", block.caption or block.alt_text or image_path.name)
        except (OSError, TypeError, ValueError, UnrecognizedImageError) as exc:
            cls._add_image_placeholder(word_document, block)
            context.warnings.append(
                cls._image_warning(block, f"image could not be embedded: {type(exc).__name__}")
            )
            return

        if block.caption:
            caption = word_document.add_paragraph(block.caption, style="Caption")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def _resolve_image(document: DocumentIR, uri: str) -> tuple[Path | None, str]:
        parsed = urlparse(uri)
        is_windows_drive_path = (
            len(parsed.scheme) == 1 and len(uri) >= 3 and uri[1:3] in {":\\", ":/"}
        )
        if parsed.scheme and not is_windows_drive_path:
            return None, "remote or embedded image URIs are not fetched"
        if parsed.fragment:
            return None, "fragment image references cannot be embedded"

        raw_path = unquote(uri if is_windows_drive_path else (parsed.path or uri))
        image_path = Path(raw_path).expanduser()
        if not image_path.is_absolute():
            source = Path(document.source.path).expanduser()
            if not source.is_absolute():
                source = source.resolve()
            image_path = source.parent / image_path
        image_path = image_path.resolve()
        if not image_path.is_file():
            return None, "image file does not exist"
        return image_path, ""

    @staticmethod
    def _add_image_placeholder(word_document: DocumentObject, block: ImageBlock) -> None:
        alt_text = block.alt_text or block.uri
        placeholder = word_document.add_paragraph(style="LDC Image Placeholder")
        placeholder.add_run(f"[Image unavailable: {alt_text}]")
        if block.caption:
            caption = word_document.add_paragraph(block.caption, style="Caption")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def _image_warning(block: ImageBlock, reason: str) -> DocumentWarning:
        return DocumentWarning(
            code="docx.image_unavailable",
            message="Image could not be embedded; alt text and caption were preserved",
            page_number=block.page_number,
            details={"reason": reason, "uri": block.uri},
        )
